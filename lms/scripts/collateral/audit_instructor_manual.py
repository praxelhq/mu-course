#!/usr/bin/env python3
"""Release audit for the consolidated Sessions 03-05 instructor manual."""

from __future__ import annotations

import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CP = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC = "http://purl.org/dc/elements/1.1/"
NS = {"w": W, "cp": CP, "dc": DC}


def attr(name: str) -> str:
    return f"{{{W}}}{name}"


def main() -> int:
    if len(sys.argv) != 2:
        raise SystemExit("usage: audit_manual.py MANUAL.docx")

    path = Path(sys.argv[1])
    failures: list[str] = []
    notes: list[str] = []

    def check(condition: bool, message: str) -> None:
        if condition:
            notes.append(f"PASS  {message}")
        else:
            failures.append(message)
            notes.append(f"FAIL  {message}")

    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")
        document = ET.fromstring(document_xml)
        styles = ET.fromstring(archive.read("word/styles.xml"))
        numbering = ET.fromstring(archive.read("word/numbering.xml"))
        core = ET.fromstring(archive.read("docProps/core.xml"))
        names = set(archive.namelist())

        sect = document.find(".//w:sectPr", NS)
        check(sect is not None, "one section-properties block is present")
        if sect is not None:
            page = sect.find("w:pgSz", NS)
            margins = sect.find("w:pgMar", NS)
            check(page is not None and page.get(attr("w")) == "12240" and page.get(attr("h")) == "15840",
                  "US Letter portrait page size is 12240 x 15840 DXA")
            check(margins is not None and all(margins.get(attr(side)) == "1440" for side in ("top", "right", "bottom", "left")),
                  "all page margins are exactly 1 inch")
            check(margins is not None and margins.get(attr("header")) == "708" and margins.get(attr("footer")) == "708",
                  "header and footer distances are 708 DXA")

        background = document.find("w:background", NS)
        check(background is not None and background.get(attr("color")) == "FBF8F3",
              "Parchment page background is encoded")

        style_ids = {node.get(attr("styleId")): node for node in styles.findall("w:style", NS)}
        check(all(style_id in style_ids for style_id in ("Normal", "Title", "Heading1", "Heading2", "Heading3")),
              "native Normal, Title, and Heading 1-3 styles exist")
        title_style = style_ids.get("Title")
        check(title_style is not None and title_style.find("w:pPr/w:pBdr", NS) is None,
              "Title style has no paragraph border")

        tables = document.findall(".//w:tbl", NS)
        check(len(tables) == 16, "document contains the expected 16 tables")
        check(all(table.find("w:tblPr/w:tblCaption", NS) is not None for table in tables),
              "every table has an accessibility caption")
        check(all(table.find("w:tblPr/w:tblDescription", NS) is not None for table in tables),
              "every table has an accessibility description")
        check(all(table.find("w:tr/w:trPr/w:tblHeader", NS) is not None for table in tables),
              "every table marks its first row as a header")
        rows = document.findall(".//w:tr", NS)
        check(all(row.find("w:trPr/w:cantSplit", NS) is not None for row in rows),
              "every table row is protected from page splitting")
        check(not document.findall(".//w:trHeight", NS), "no fixed table-row heights are present")

        num_paragraphs = document.findall(".//w:pPr/w:numPr", NS)
        num_ids = {node.get(attr("val")) for node in document.findall(".//w:pPr/w:numPr/w:numId", NS)}
        abstract_ids = numbering.findall("w:abstractNum", NS)
        check(len(num_paragraphs) >= 150, "real Word list numbering is used extensively")
        check(len(num_ids) >= 7 and len(abstract_ids) >= 3,
              "multiple list instances preserve bullets, checklists, and numbered-list restarts")

        core_values = []
        for tag in (f"{{{DC}}}creator", f"{{{CP}}}lastModifiedBy"):
            node = core.find(tag)
            core_values.append((node.text or "").strip() if node is not None else "")
        check(not any(core_values), "creator and last-modified-by metadata are blank")
        check("docProps/custom.xml" not in names, "custom document properties are absent")
        check(b"rsid" not in document_xml.lower(), "revision-session identifiers are absent")

    doc = Document(path)
    full_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    full_text += "\n" + "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    required_phrases = {
        "lifecycle truth": "All three source packages are Authored",
        "automatic V2": "System automatically grants one immutable V2",
        "receipt-timestamp expiry": "ordinary expiry is 10 calendar days from that learner's recorded V1-receipt timestamp",
        "no silent due-date redefinition": "Changing a personal V1 due date must not silently redefine an already-created V2 grant",
        "audited expiry change": "unless the audited extension explicitly changes that grant's expiry",
        "V2 limit": "never create V3 or a second V2 grant",
        "blueprint privacy": "Raw blueprint remains private",
        "roster gate": "blueprint remains roster-gated private evidence even after scrubbing",
        "nomination boundary": "a team may nominate, but the instructor alone selects exactly one existing finalised individual submission version",
        "selection authority": "do not merge an integrated package or delegate selection authority to the team",
    }
    for label, phrase in required_phrases.items():
        check(phrase.casefold() in full_text.casefold(), f"required content control is present: {label}")

    stale_expiry_phrases = (
        "10 calendar days after section V1 deadline",
        "10 calendar days after that personal V1 deadline",
    )
    check(not any(phrase.casefold() in full_text.casefold() for phrase in stale_expiry_phrases),
          "no stale deadline-anchored V2 expiry wording remains")

    run_tables = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers == ["Time", "Min", "Mode", "Instructor move", "Learner work / evidence"]:
            minutes = [int(row.cells[1].text.strip()) for row in table.rows[1:]]
            run_tables.append(sum(minutes))
    check(run_tables == [120, 120, 120], "all three run-of-show tables sum exactly to 120 minutes")

    for line in notes:
        print(line)
    print(f"SUMMARY: {len(notes) - len(failures)} passed; {len(failures)} failed")
    return 1 if failures else 0


if __name__ == "__main__":
    raise SystemExit(main())
