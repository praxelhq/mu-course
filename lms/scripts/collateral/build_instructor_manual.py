#!/usr/bin/env python3
"""Build the Sessions 03-05 instructor manual from its Markdown source."""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


# Brand palette: named overrides on compact_reference_guide.
PARCHMENT = "FBF8F3"
INK = "1F1A14"
PINE = "1E3A35"
OCHRE = "C4581A"
BEACON = "F0D478"
SAND = "EDE5D8"
CHARCOAL = "5C5046"
CLAY = "9C8E82"
CREAM = "F5F0E8"
NEAR_WHITE = "FDFBF8"
RISK = "8A2D22"

BODY_FONT = "Arial"
DISPLAY_FONT = "Georgia"
MONO_FONT = "Menlo"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120

INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
NUMBERED_RE = re.compile(r"^(\d+)\.\s+(.*)$")


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_font(run, name: str, size: float | None = None, color: str | None = None,
             bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name: str, size: float, color: str = INK,
                   bold: bool = False, italic: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    style.font.italic = italic
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_paragraph_tokens(style, before: float, after: float, line_spacing: float,
                         keep_with_next: bool = False) -> None:
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing
    pf.widow_control = True
    pf.keep_with_next = keep_with_next


def set_style_shading(style, fill: str) -> None:
    ppr = style.element.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_style_left_border(style, color: str, size: int = 16, space: int = 8) -> None:
    ppr = style.element.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    pbdr.append(left)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = CELL_TOP_BOTTOM_DXA,
                     bottom: int = CELL_TOP_BOTTOM_DXA,
                     start: int = CELL_START_END_DXA,
                     end: int = CELL_START_END_DXA) -> None:
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = SAND, size: int = 6) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblpr = tbl.tblPr

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent))
    tblind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = trpr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        trpr.append(header)
    header.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    cant = trpr.find(qn("w:cantSplit"))
    if cant is None:
        cant = OxmlElement("w:cantSplit")
        trpr.append(cant)


def set_table_accessibility(table, caption: str, description: str) -> None:
    tblpr = table._tbl.tblPr
    for tag, value in (("tblCaption", caption), ("tblDescription", description)):
        node = tblpr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tblpr.append(node)
        node.set(qn("w:val"), value[:255])


def add_hyperlink(paragraph, text: str, url: str, color: str = PINE) -> None:
    # Only permit ordinary web links in the generated document.
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        paragraph.add_run(text)
        return
    rid = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), BODY_FONT)
    rpr.append(rfonts)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    rpr.append(color_el)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, source: str, base_size: float | None = None,
               base_color: str | None = None) -> None:
    pos = 0
    for match in INLINE_RE.finditer(source):
        if match.start() > pos:
            run = paragraph.add_run(source[pos:match.start()])
            if base_size or base_color:
                set_font(run, BODY_FONT, base_size, base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, BODY_FONT, base_size, base_color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, MONO_FONT, (base_size or 10.5) - 0.4, PINE)
            rpr = run._element.get_or_add_rPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:fill"), CREAM)
            rpr.append(shd)
        else:
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        pos = match.end()
    if pos < len(source):
        run = paragraph.add_run(source[pos:])
        if base_size or base_color:
            set_font(run, BODY_FONT, base_size, base_color)


def parse_front_matter(lines: list[str]) -> tuple[dict[str, str], list[str]]:
    if not lines or lines[0].strip() != "---":
        return {}, lines
    metadata: dict[str, str] = {}
    end = None
    for idx in range(1, len(lines)):
        if lines[idx].strip() == "---":
            end = idx
            break
        if ":" in lines[idx]:
            key, value = lines[idx].split(":", 1)
            metadata[key.strip()] = value.strip()
    if end is None:
        raise ValueError("Unclosed Markdown front matter")
    return metadata, lines[end + 1:]


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def column_widths(headers: list[str]) -> list[int]:
    n = len(headers)
    lowered = [h.lower() for h in headers]
    if n == 5 and lowered[:3] == ["time", "min", "mode"]:
        return [1050, 600, 1450, 3100, 3160]
    if n == 4 and lowered[0] == "session" and "must pass" in lowered[-1]:
        return [950, 2100, 1900, 4410]
    if n == 4 and lowered[-1].startswith("recheck / operator"):
        return [2100, 3000, 1300, 2960]
    if n == 4 and "operator record" in lowered[-1]:
        return [2300, 3000, 2500, 1560]
    if n == 4 and lowered[0] in {"session", "label"}:
        return [1700, 3100, 2950, 1610]
    if n == 4 and "operator action" in lowered[-1]:
        return [1550, 2200, 2800, 2810]
    if n == 4:
        return [1950, 2700, 2700, 2010]
    if n == 3:
        return [2100, 3500, 3760]
    if n == 2:
        return [2250, 7110]
    if n == 5:
        return [1500, 1650, 1900, 2200, 2110]
    base = CONTENT_WIDTH_DXA // n
    widths = [base] * n
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_abstract_numbering(doc: Document, fmt: str, text_value: str,
                           left: int = 540, hanging: int = 271) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), fmt)
    lvl.append(numfmt)
    lvltext = OxmlElement("w:lvlText")
    lvltext.set(qn("w:val"), text_value)
    lvl.append(lvltext)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    marker_font = "Arial Unicode MS" if text_value == "☐" else BODY_FONT
    for attr in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), marker_font)
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), PINE)
    rpr.append(color)
    lvl.append(rpr)
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is not None:
        numbering.insert(numbering.index(first_num), abstract)
    else:
        numbering.append(abstract)
    return abstract_id


def instantiate_numbering(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    num_id = (max(num_ids) + 1) if num_ids else 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    # Force each instantiated sequence to start at 1. LibreOffice otherwise
    # continues numbering across distinct lists that share one abstractNum.
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)


def add_field(paragraph, instruction: str, fallback: str = "1") -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_font(run, BODY_FONT, 8.5, CHARCOAL)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, BODY_FONT, 11, INK)
    set_paragraph_tokens(normal, 0, 6, 1.25)

    title = styles["Title"]
    set_style_font(title, DISPLAY_FONT, 30, PINE, bold=True)
    set_paragraph_tokens(title, 0, 8, 1.0, keep_with_next=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, BODY_FONT, 13.5, CHARCOAL)
    set_paragraph_tokens(subtitle, 0, 18, 1.15, keep_with_next=True)

    # Remove Word-template border residue from the opening title stack.
    for clean_style in (title, subtitle):
        ppr = clean_style.element.get_or_add_pPr()
        pbdr = ppr.find(qn("w:pBdr"))
        if pbdr is not None:
            ppr.remove(pbdr)

    heading_tokens = {
        "Heading 1": (DISPLAY_FONT, 16, PINE, 18, 10),
        "Heading 2": (BODY_FONT, 13, PINE, 14, 7),
        "Heading 3": (BODY_FONT, 12, CHARCOAL, 10, 5),
    }
    for name, (font, size, color, before, after) in heading_tokens.items():
        style = styles[name]
        set_style_font(style, font, size, color, bold=True)
        set_paragraph_tokens(style, before, after, 1.10, keep_with_next=True)
        style.paragraph_format.keep_together = True
        ppr = style.element.get_or_add_pPr()
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            ppr.append(outline)
        outline.set(qn("w:val"), str(int(name[-1]) - 1))

    def add_style(name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
        return styles[name] if name in styles else styles.add_style(name, style_type)

    kicker = add_style("MU Kicker")
    set_style_font(kicker, BODY_FONT, 8.5, OCHRE, bold=True)
    set_paragraph_tokens(kicker, 0, 6, 1.0, keep_with_next=True)
    rpr = kicker.element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "14")
    rpr.append(spacing)

    status = add_style("MU Status")
    set_style_font(status, BODY_FONT, 9.2, CHARCOAL, italic=True)
    set_paragraph_tokens(status, 0, 16, 1.15)

    lead = add_style("MU Lead")
    set_style_font(lead, BODY_FONT, 10.5, PINE, bold=False)
    set_paragraph_tokens(lead, 6, 8, 1.20, keep_with_next=True)

    source = add_style("MU Source Note")
    set_style_font(source, BODY_FONT, 9, CHARCOAL, italic=True)
    set_paragraph_tokens(source, 6, 8, 1.15)

    quote = add_style("MU Quote")
    set_style_font(quote, BODY_FONT, 10.5, INK)
    set_paragraph_tokens(quote, 6, 10, 1.20)
    quote.paragraph_format.left_indent = Inches(0.16)
    quote.paragraph_format.right_indent = Inches(0.08)
    set_style_shading(quote, CREAM)
    set_style_left_border(quote, OCHRE)

    for name in ("MU Bullet", "MU Number", "MU Checklist"):
        style = add_style(name)
        set_style_font(style, BODY_FONT, 11, INK)
        set_paragraph_tokens(style, 0, 4, 1.25)

    table_header = add_style("MU Table Header")
    set_style_font(table_header, BODY_FONT, 8.6, CREAM, bold=True)
    set_paragraph_tokens(table_header, 0, 0, 1.05)

    table_body = add_style("MU Table Body")
    set_style_font(table_body, BODY_FONT, 9.0, INK)
    set_paragraph_tokens(table_body, 0, 0, 1.10)

    table_spacer = add_style("MU Table Spacer")
    set_style_font(table_spacer, BODY_FONT, 2, PARCHMENT)
    set_paragraph_tokens(table_spacer, 0, 3, 1.0)

    metric_number = add_style("MU Metric Number")
    set_style_font(metric_number, DISPLAY_FONT, 19, PINE, bold=True)
    set_paragraph_tokens(metric_number, 0, 1, 1.0)
    metric_number.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    metric_label = add_style("MU Metric Label")
    set_style_font(metric_label, BODY_FONT, 8, CHARCOAL, bold=True)
    set_paragraph_tokens(metric_label, 0, 0, 1.0)
    metric_label.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header = add_style("MU Header")
    set_style_font(header, BODY_FONT, 8.2, CHARCOAL, bold=True)
    set_paragraph_tokens(header, 0, 0, 1.0)

    footer = add_style("MU Footer")
    set_style_font(footer, BODY_FONT, 8.2, CHARCOAL)
    set_paragraph_tokens(footer, 0, 0, 1.0)

    # A real table style, supplemented with explicit cell/table XML for portability.
    table_style = add_style("MU Operator Table", WD_STYLE_TYPE.TABLE)
    set_style_font(table_style, BODY_FONT, 9, INK)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    # Very light page tone. This is subtle enough for print and keeps brand continuity.
    root = doc._element
    old_bg = root.find(qn("w:background"))
    if old_bg is not None:
        root.remove(old_bg)
    background = OxmlElement("w:background")
    background.set(qn("w:color"), PARCHMENT)
    root.insert(0, background)

    settings = doc.settings.element
    display = settings.find(qn("w:displayBackgroundShape"))
    if display is None:
        display = OxmlElement("w:displayBackgroundShape")
        settings.append(display)
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    header = section.header
    hp = header.paragraphs[0]
    hp.style = doc.styles["MU Header"]
    hp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    hp.add_run("SESSIONS 03-05  /  INSTRUCTOR MANUAL")
    hp.add_run("\tAUTHORED SOURCE")

    fp = section.footer.paragraphs[0]
    fp.style = doc.styles["MU Footer"]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    fp.add_run("INSTRUCTOR / EVALUATOR USE ONLY")
    fp.add_run("\t")
    add_field(fp, "PAGE", "1")
    fp.add_run(" / ")
    add_field(fp, "NUMPAGES", "1")


def add_cover(doc: Document, meta: dict[str, str]) -> None:
    p = doc.add_paragraph(style="MU Kicker")
    p.add_run("INSTRUCTOR OPERATIONS  /  COURSE 1")

    p = doc.add_paragraph(style="Title")
    p.add_run(meta.get("title", "Sessions 03-05 Instructor Manual"))

    p = doc.add_paragraph(style="Subtitle")
    p.add_run(meta.get("subtitle", "Data evidence, a working app, and a controlled revenue system"))

    metrics = [("03", "DATA EVIDENCE"), ("04", "WORKING APP"), ("05", "REVENUE SYSTEM"), ("120", "MIN EACH")]
    table = doc.add_table(rows=1, cols=4)
    table.style = "MU Operator Table"
    set_table_geometry(table, [2340, 2340, 2340, 2340], indent=100)
    set_table_borders(table, SAND, 6)
    set_table_accessibility(table, "Session sequence metric strip",
                            "Four cells identify Sessions 03, 04, 05, and the 120-minute duration for each.")
    for idx, (number, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, CREAM if idx < 3 else PARCHMENT)
        set_cell_margins(cell, top=120, bottom=120, start=100, end=100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = cell.paragraphs[0]
        p1.style = "MU Metric Number"
        p1.add_run(number)
        p2 = cell.add_paragraph(style="MU Metric Label")
        p2.add_run(label)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])

    p = doc.add_paragraph(style="MU Status")
    p.add_run(f"Version {meta.get('version', '1.0')}  |  Prepared {meta.get('prepared', '30 July 2026')}  |  ")
    p.add_run(meta.get("classification", "Instructor and evaluator use only"))
    p.add_run("\n")
    p.add_run(meta.get("source_status", "Authored packages; validation and rehearsal pending"))


def add_table(doc: Document, rows: list[list[str]], context: str, table_index: int) -> None:
    if len(rows) < 2:
        return
    headers = rows[0]
    col_count = len(headers)
    if any(len(row) != col_count for row in rows):
        raise ValueError(f"Inconsistent table columns under {context}: {rows}")
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "MU Operator Table"
    widths = column_widths(headers)
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_table_accessibility(table, f"Table {table_index}: {context}",
                            f"Accessible reference table under heading {context}; first row is the repeating header.")
    set_repeat_table_header(table.rows[0])

    centered_headers = {"time", "min", "score", "version", "session"}
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        set_row_cant_split(row)
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_shading(cell, PINE if r_idx == 0 else (NEAR_WHITE if r_idx % 2 else CREAM))
            p = cell.paragraphs[0]
            p.style = "MU Table Header" if r_idx == 0 else "MU Table Body"
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                           if headers[c_idx].strip().lower() in centered_headers
                           else WD_ALIGN_PARAGRAPH.LEFT)
            add_inline(p, value, base_size=8.6 if r_idx == 0 else 9.0,
                       base_color=CREAM if r_idx == 0 else INK)
    doc.add_paragraph(style="MU Table Spacer")


def build(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    metadata, lines = parse_front_matter(lines)

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = metadata.get("title", "Sessions 03-05 Instructor Manual")
    doc.core_properties.subject = "Instructor operations manual for Sessions 03-05"
    doc.core_properties.keywords = "instructor manual, session run-of-show, grading, LMS gates, fallbacks"
    doc.core_properties.category = "Course operations"
    doc.core_properties.comments = "Generated from the consolidated Markdown source."
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""

    abstract_bullet = add_abstract_numbering(doc, "bullet", "•")
    abstract_check = add_abstract_numbering(doc, "bullet", "☐")
    abstract_decimal = add_abstract_numbering(doc, "decimal", "%1.")
    bullet_num = instantiate_numbering(doc, abstract_bullet)
    check_num = instantiate_numbering(doc, abstract_check)

    add_cover(doc, metadata)

    # Skip the Markdown title because the cover already owns it.
    skipped_title = False
    body_h1_seen = False
    idx = 0
    table_index = 0
    current_heading = metadata.get("title", "Instructor manual")
    previous_kind = "cover"
    active_decimal_num: int | None = None

    while idx < len(lines):
        raw = lines[idx].rstrip()
        stripped = raw.strip()

        if not stripped:
            previous_kind = "blank"
            active_decimal_num = None
            idx += 1
            continue

        if stripped == "<!-- PAGEBREAK -->":
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)
            previous_kind = "pagebreak"
            active_decimal_num = None
            idx += 1
            continue

        if stripped.startswith("#"):
            match = re.match(r"^(#{1,4})\s+(.*)$", stripped)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                if level == 1 and not skipped_title and text == metadata.get("title"):
                    skipped_title = True
                    idx += 1
                    continue
                # Before the first session H1, the Markdown title is skipped;
                # promote its H2/H3/H4 children by one level in Word so the
                # navigational hierarchy starts at Heading 1 without a skip.
                if level == 1:
                    style_level = 1
                    body_h1_seen = True
                elif not body_h1_seen:
                    style_level = level - 1
                else:
                    style_level = level
                style_level = min(max(style_level, 1), 3)
                style = f"Heading {style_level}"
                p = doc.add_paragraph(style=style)
                add_inline(p, text)
                current_heading = text
                previous_kind = "heading"
                active_decimal_num = None
                idx += 1
                continue

        if stripped.startswith("|") and idx + 1 < len(lines) and TABLE_SEPARATOR_RE.match(lines[idx + 1]):
            table_rows = [parse_table_row(stripped)]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_rows.append(parse_table_row(lines[idx]))
                idx += 1
            table_index += 1
            add_table(doc, table_rows, current_heading, table_index)
            previous_kind = "table"
            active_decimal_num = None
            continue

        if stripped.startswith(">"):
            text = stripped[1:].strip()
            p = doc.add_paragraph(style="MU Quote")
            add_inline(p, text)
            previous_kind = "quote"
            active_decimal_num = None
            idx += 1
            continue

        if stripped.startswith("- "):
            text = stripped[2:].strip()
            is_check = text.startswith("[ ]")
            if is_check:
                text = text[3:].strip()
            p = doc.add_paragraph(style="MU Checklist" if is_check else "MU Bullet")
            apply_num(p, check_num if is_check else bullet_num)
            add_inline(p, text)
            previous_kind = "check" if is_check else "bullet"
            active_decimal_num = None
            idx += 1
            continue

        numbered = NUMBERED_RE.match(stripped)
        if numbered:
            if previous_kind != "number" or active_decimal_num is None:
                active_decimal_num = instantiate_numbering(doc, abstract_decimal)
            p = doc.add_paragraph(style="MU Number")
            apply_num(p, active_decimal_num)
            add_inline(p, numbered.group(2))
            previous_kind = "number"
            idx += 1
            continue

        style = "Normal"
        if stripped.startswith("**Source binding:"):
            style = "MU Source Note"
        elif (stripped.startswith("**Protected student work:")
              or stripped.startswith("**Timing provenance:")
              or stripped.startswith("**No-Go wording:")):
            style = "MU Lead"
        p = doc.add_paragraph(style=style)
        add_inline(p, stripped)
        previous_kind = "paragraph"
        active_decimal_num = None
        idx += 1

    # Keep generated OOXML free of editing-session identifiers and personal metadata.
    for element in doc.element.iter():
        for attr_name in list(element.attrib):
            if attr_name.rsplit("}", 1)[-1].lower().startswith("rsid"):
                del element.attrib[attr_name]
        for child in list(element):
            if child.tag.rsplit("}", 1)[-1].lower().startswith("rsid"):
                element.remove(child)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build(args.markdown, args.output)


if __name__ == "__main__":
    main()
